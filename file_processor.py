"""
File Processor Module
Extracts text from various file formats for PII anonymization.

Supported formats:
- Plain text (.txt)
- PDF (.pdf)
- Word documents (.docx)
- Markdown (.md)
- CSV (.csv)
- JSON (.json)
"""

import os
from pathlib import Path
from typing import Tuple, Optional, Union
from abc import ABC, abstractmethod
import json
import csv
import io
import hashlib


class FileProcessor(ABC):
    """Base class for file processors."""
    
    @abstractmethod
    def extract_text(self, content: bytes) -> str:
        """Extract text content from file bytes."""
        pass
    
    @abstractmethod
    def reconstruct(self, original_content: bytes, anonymized_text: str) -> bytes:
        """Reconstruct file with anonymized text."""
        pass
    
    @staticmethod
    def get_file_hash(content: bytes) -> str:
        """Get SHA-256 hash of file content."""
        return hashlib.sha256(content).hexdigest()


class TextProcessor(FileProcessor):
    """Process plain text files."""
    
    def extract_text(self, content: bytes) -> str:
        # Try different encodings
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode text file with any known encoding")
    
    def reconstruct(self, original_content: bytes, anonymized_text: str) -> bytes:
        return anonymized_text.encode('utf-8')


class PDFProcessor(FileProcessor):
    """Process PDF files using pypdf."""
    
    def __init__(self):
        try:
            import pypdf
            self.pypdf = pypdf
        except ImportError:
            raise ImportError("pypdf required: pip install pypdf")
    
    def extract_text(self, content: bytes) -> str:
        reader = self.pypdf.PdfReader(io.BytesIO(content))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    
    def reconstruct(self, original_content: bytes, anonymized_text: str) -> bytes:
        """
        For PDFs, we return the anonymized text as a new text file.
        Full PDF reconstruction would require preserving layout which is complex.
        """
        # TODO: For production, use reportlab or similar to create new PDF
        return anonymized_text.encode('utf-8')


class DocxProcessor(FileProcessor):
    """Process Word documents using python-docx."""
    
    def __init__(self):
        try:
            import docx
            self.docx = docx
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")
    
    def extract_text(self, content: bytes) -> str:
        doc = self.docx.Document(io.BytesIO(content))
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    
    def reconstruct(self, original_content: bytes, anonymized_text: str) -> bytes:
        """
        Create new DOCX with anonymized text.
        Preserves basic structure but not complex formatting.
        """
        doc = self.docx.Document()
        for para in anonymized_text.split("\n"):
            if para.strip():
                doc.add_paragraph(para)
        
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()


class MarkdownProcessor(TextProcessor):
    """Process Markdown files (extends text processor)."""
    pass


class CSVProcessor(FileProcessor):
    """Process CSV files, anonymizing cell contents."""
    
    def extract_text(self, content: bytes) -> str:
        text = content.decode('utf-8')
        # Return as-is for text extraction, preserve structure
        return text
    
    def reconstruct(self, original_content: bytes, anonymized_text: str) -> bytes:
        return anonymized_text.encode('utf-8')


class JSONProcessor(FileProcessor):
    """Process JSON files, extracting string values."""
    
    def extract_text(self, content: bytes) -> str:
        data = json.loads(content.decode('utf-8'))
        # Extract all string values recursively
        strings = []
        self._extract_strings(data, strings)
        return "\n".join(strings)
    
    def _extract_strings(self, obj, strings: list, path: str = ""):
        if isinstance(obj, str):
            strings.append(f"{path}: {obj}")
        elif isinstance(obj, dict):
            for k, v in obj.items():
                self._extract_strings(v, strings, f"{path}.{k}" if path else k)
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                self._extract_strings(item, strings, f"{path}[{i}]")
    
    def reconstruct(self, original_content: bytes, anonymized_text: str) -> bytes:
        # For JSON, we need to replace values in the original structure
        # This is a simplified version - production would need smarter replacement
        return anonymized_text.encode('utf-8')


class FileProcessorFactory:
    """Factory for getting appropriate file processor."""
    
    PROCESSORS = {
        '.txt': TextProcessor,
        '.text': TextProcessor,
        '.pdf': PDFProcessor,
        '.docx': DocxProcessor,
        '.doc': DocxProcessor,  # Note: .doc requires different handling
        '.md': MarkdownProcessor,
        '.markdown': MarkdownProcessor,
        '.csv': CSVProcessor,
        '.json': JSONProcessor,
    }
    
    @classmethod
    def get_processor(cls, file_path: Union[str, Path]) -> FileProcessor:
        """Get appropriate processor for file type."""
        ext = Path(file_path).suffix.lower()
        
        if ext not in cls.PROCESSORS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {list(cls.PROCESSORS.keys())}")
        
        return cls.PROCESSORS[ext]()
    
    @classmethod
    def get_processor_by_type(cls, file_type: str) -> FileProcessor:
        """Get processor by file extension string."""
        if not file_type.startswith('.'):
            file_type = '.' + file_type
        return cls.get_processor(f"dummy{file_type}")
    
    @classmethod
    def is_supported(cls, file_path: Union[str, Path]) -> bool:
        """Check if file type is supported."""
        ext = Path(file_path).suffix.lower()
        return ext in cls.PROCESSORS
    
    @classmethod
    def supported_types(cls) -> list:
        """Get list of supported file extensions."""
        return list(cls.PROCESSORS.keys())


def process_file(file_path: Union[str, Path]) -> Tuple[str, str, bytes]:
    """
    Read and extract text from a file.
    
    Returns: (text_content, file_type, original_bytes)
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    
    with open(path, 'rb') as f:
        content = f.read()
    
    processor = FileProcessorFactory.get_processor(path)
    text = processor.extract_text(content)
    
    return text, path.suffix.lower(), content


def save_anonymized_file(
    original_path: Union[str, Path],
    anonymized_text: str,
    output_path: Optional[Union[str, Path]] = None
) -> Path:
    """
    Save anonymized content to a new file.
    
    If output_path is not provided, creates a new file with '_anonymized' suffix.
    """
    original = Path(original_path)
    
    if output_path is None:
        output_path = original.parent / f"{original.stem}_anonymized{original.suffix}"
    
    output = Path(output_path)
    
    # Read original content
    with open(original, 'rb') as f:
        original_content = f.read()
    
    # Get processor and reconstruct
    processor = FileProcessorFactory.get_processor(original)
    anonymized_content = processor.reconstruct(original_content, anonymized_text)
    
    # Write output
    with open(output, 'wb') as f:
        f.write(anonymized_content)
    
    return output


if __name__ == "__main__":
    # Test with a sample text file
    import tempfile
    
    print("=== File Processor Tests ===\n")
    
    # Test text file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        f.write("Patient John Smith, SSN 123-45-6789, called from 555-123-4567.")
        temp_path = f.name
    
    try:
        text, file_type, content = process_file(temp_path)
        print(f"File type: {file_type}")
        print(f"Extracted text: {text}")
        print(f"File size: {len(content)} bytes")
    finally:
        os.unlink(temp_path)
    
    print("\n=== Supported Types ===")
    print(FileProcessorFactory.supported_types())
